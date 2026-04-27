import os
import re
from typing import Optional


class DocumentParser:
    def parse_pdf(self, file_path: str) -> str:
        import fitz
        doc = fitz.open(file_path)
        text_parts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(f"[第{page_num + 1}页]\n{text}")
            else:
                try:
                    import pytesseract
                    from PIL import Image
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                    if ocr_text.strip():
                        text_parts.append(f"[第{page_num + 1}页 - OCR]\n{ocr_text}")
                except Exception:
                    pass
        doc.close()
        return "\n\n".join(text_parts)

    def parse_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(f"[表格] {row_text}")
        return "\n\n".join(text_parts)

    def parse_file(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        return text.strip()
